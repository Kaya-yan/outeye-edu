"""
文档解析器

支持多种文档格式的解析和文本分块
"""

from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path
import re
import hashlib


@dataclass
class DocumentChunk:
    """文档块"""
    id: str
    doc_id: str
    content: str
    metadata: Dict[str, any] = field(default_factory=dict)
    start_pos: int = 0
    end_pos: int = 0
    chunk_index: int = 0


@dataclass
class ParsedDocument:
    """解析后的文档"""
    id: str
    title: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, any] = field(default_factory=dict)
    file_path: str = ""
    file_type: str = ""


class DocumentParser:
    """文档解析器"""

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        min_chunk_size: int = 100
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size

    def parse_file(self, file_path: str) -> ParsedDocument:
        """
        解析文件

        Args:
            file_path: 文件路径

        Returns:
            解析后的文档
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 根据文件类型选择解析方法
        suffix = path.suffix.lower()

        if suffix == '.pdf':
            content, metadata = self._parse_pdf(path)
        elif suffix in ['.docx', '.doc']:
            content, metadata = self._parse_docx(path)
        elif suffix == '.md':
            content, metadata = self._parse_markdown(path)
        elif suffix == '.txt':
            content, metadata = self._parse_text(path)
        elif suffix in ['.html', '.htm']:
            content, metadata = self._parse_html(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

        # 生成文档ID
        doc_id = self._generate_doc_id(file_path)

        # 提取标题
        title = self._extract_title(content, path.name)

        # 文本分块
        chunks = self._chunk_text(content, doc_id)

        return ParsedDocument(
            id=doc_id,
            title=title,
            content=content,
            chunks=chunks,
            metadata=metadata,
            file_path=str(path),
            file_type=suffix
        )

    def parse_text(
        self,
        text: str,
        title: str = "",
        metadata: Dict[str, any] = None
    ) -> ParsedDocument:
        """
        解析文本

        Args:
            text: 文本内容
            title: 标题
            metadata: 元数据

        Returns:
            解析后的文档
        """
        # 生成文档ID
        doc_id = self._generate_doc_id(text[:100])

        # 文本分块
        chunks = self._chunk_text(text, doc_id)

        return ParsedDocument(
            id=doc_id,
            title=title or "未命名文档",
            content=text,
            chunks=chunks,
            metadata=metadata or {}
        )

    def _parse_pdf(self, path: Path) -> Tuple[str, Dict]:
        """解析PDF文件"""
        try:
            import PyPDF2

            content = ""
            metadata = {}

            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                # 提取元数据
                if reader.metadata:
                    metadata = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', '')
                    }

                # 提取文本
                for page in reader.pages:
                    content += page.extract_text() + "\n"

            return content.strip(), metadata

        except ImportError:
            # 如果没有PyPDF2，尝试使用pdfplumber
            try:
                import pdfplumber

                content = ""
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"

                return content.strip(), {}
            except ImportError:
                raise ImportError("需要安装PyPDF2或pdfplumber来解析PDF文件")

    def _parse_docx(self, path: Path) -> Tuple[str, Dict]:
        """解析Word文档"""
        try:
            from docx import Document

            doc = Document(path)

            # 提取元数据
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else ''
            }

            # 提取文本
            content = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"

            return content.strip(), metadata

        except ImportError:
            # 如果没有python-docx，尝试使用unzip
            import zipfile
            import xml.etree.ElementTree as ET

            content = ""
            metadata = {}

            with zipfile.ZipFile(path, 'r') as z:
                # 读取document.xml
                with z.open('word/document.xml') as f:
                    tree = ET.parse(f)
                    root = tree.getroot()

                    # 提取文本
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    for elem in root.findall('.//w:t', ns):
                        if elem.text:
                            content += elem.text

            return content.strip(), metadata

    def _parse_markdown(self, path: Path) -> Tuple[str, Dict]:
        """解析Markdown文件"""
        content = path.read_text(encoding='utf-8')

        # 提取YAML frontmatter
        metadata = {}
        frontmatter_match = re.match(r'^---\n(.*?)\n---\n', content, re.DOTALL)
        if frontmatter_match:
            try:
                import yaml
                metadata = yaml.safe_load(frontmatter_match.group(1))
            except:
                pass
            content = content[frontmatter_match.end():]

        # 移除Markdown格式
        # 保留文本内容，移除格式标记
        clean_content = content
        clean_content = re.sub(r'^#{1,6}\s+', '', clean_content, flags=re.MULTILINE)  # 标题
        clean_content = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_content)  # 粗体
        clean_content = re.sub(r'\*(.*?)\*', r'\1', clean_content)  # 斜体
        clean_content = re.sub(r'`(.*?)`', r'\1', clean_content)  # 行内代码
        clean_content = re.sub(r'```[\s\S]*?```', '', clean_content)  # 代码块
        clean_content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_content)  # 链接
        clean_content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', clean_content)  # 图片

        return clean_content.strip(), metadata

    def _parse_text(self, path: Path) -> Tuple[str, Dict]:
        """解析纯文本文件"""
        content = path.read_text(encoding='utf-8')
        return content.strip(), {}

    def _parse_html(self, path: Path) -> Tuple[str, Dict]:
        """解析HTML文件"""
        try:
            from bs4 import BeautifulSoup

            html_content = path.read_text(encoding='utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')

            # 提取元数据
            metadata = {}
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.text

            # 提取文本
            content = soup.get_text(separator='\n', strip=True)

            return content, metadata

        except ImportError:
            # 简单的HTML解析
            html_content = path.read_text(encoding='utf-8')

            # 移除HTML标签
            content = re.sub(r'<[^>]+>', ' ', html_content)
            content = re.sub(r'\s+', ' ', content)

            return content.strip(), {}

    def _generate_doc_id(self, content: str) -> str:
        """生成文档ID"""
        return hashlib.md5(content.encode()).hexdigest()[:16]

    def _extract_title(self, content: str, filename: str) -> str:
        """提取标题"""
        # 尝试从内容中提取标题
        lines = content.split('\n')
        for line in lines[:10]:  # 检查前10行
            line = line.strip()
            if line and len(line) < 100:
                # 如果是Markdown标题
                if line.startswith('#'):
                    return line.lstrip('#').strip()
                # 如果是短行，可能是标题
                if len(line) < 50:
                    return line

        # 使用文件名
        return Path(filename).stem

    def _chunk_text(self, text: str, doc_id: str) -> List[DocumentChunk]:
        """
        文本分块

        Args:
            text: 文本内容
            doc_id: 文档ID

        Returns:
            文档块列表
        """
        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            # 计算结束位置
            end = min(start + self.chunk_size, len(text))

            # 尝试在句子边界分割
            if end < len(text):
                # 寻找最近的句号、问号、感叹号
                last_period = max(
                    text.rfind('.', start, end),
                    text.rfind('?', start, end),
                    text.rfind('!', start, end),
                    text.rfind('。', start, end),
                    text.rfind('？', start, end),
                    text.rfind('！', start, end)
                )

                if last_period > start + self.min_chunk_size:
                    end = last_period + 1

            # 提取块内容
            chunk_content = text[start:end].strip()

            # 跳过太短的块
            if len(chunk_content) < self.min_chunk_size:
                start = end - self.chunk_overlap
                if start >= len(text):
                    break
                continue

            # 生成块ID
            chunk_id = f"{doc_id}_chunk_{chunk_index}"

            # 创建块
            chunk = DocumentChunk(
                id=chunk_id,
                doc_id=doc_id,
                content=chunk_content,
                start_pos=start,
                end_pos=end,
                chunk_index=chunk_index
            )

            chunks.append(chunk)

            # 移动到下一个块（考虑重叠）
            start = end - self.chunk_overlap
            chunk_index += 1

            # 防止无限循环
            if start >= len(text):
                break

        return chunks

    def merge_chunks(self, chunks: List[DocumentChunk], window: int = 3) -> List[str]:
        """
        合并相邻块以提供更完整的上下文

        Args:
            chunks: 文档块列表
            window: 合并窗口大小

        Returns:
            合并后的文本列表
        """
        merged = []

        for i in range(len(chunks)):
            # 获取窗口内的块
            start_idx = max(0, i - window // 2)
            end_idx = min(len(chunks), i + window // 2 + 1)

            window_chunks = chunks[start_idx:end_idx]
            merged_text = " ".join(c.content for c in window_chunks)

            merged.append(merged_text)

        return merged
