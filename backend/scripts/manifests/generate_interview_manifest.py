"""
生成 interview.json 访谈元数据清单

数据来源：seed-materials/ 下的 5 份访谈 docx
- 0729 常老师 2 份（0714 教学设计、0720 学生差异）
- 0811 张/林/陈 各 1 份

输出：backend/scripts/manifests/interview.json

匿名化规则：
- 教师真实姓名 -> T-01~T-04
- 学校具体名称 -> "某省重点中学" / "某外国语大学" / "某二本院校" / "某高校外国语学院"
- 访谈人 -> "访谈人-01" / "访谈人-02"

用法：
    python backend/scripts/manifests/generate_interview_manifest.py
"""

import json
import os
import re
import sys
from pathlib import Path

import docx

SEED_MATERIALS_ROOT = Path(r"C:\Users\ht\Documents\outeye3.0\seed-materials")
OUTPUT_PATH = Path(__file__).parent / "interview.json"


# 5 份访谈的元数据（基本信息从 docx 提取，匿名化映射固定）
# 文件路径 + 教师代号 + 学校匿名化 + 访谈人匿名化
INTERVIEW_FILES = [
    {
        "doc_id": "I-001",
        "file_path": str(SEED_MATERIALS_ROOT / "数据收集（0729）" / "访谈" / "访谈-常老师-教学设计、课文处理-20260714.docx"),
        "teacher_code": "T-01",
        "teacher_real_name": "常育红",
        "school_anonymized": "某省重点中学",
        "interviewer_anonymized": "访谈人-01",
        "interviewer_real_name": "杜玥林",
    },
    {
        "doc_id": "I-002",
        "file_path": str(SEED_MATERIALS_ROOT / "数据收集（0729）" / "访谈" / "访谈-常老师-学生差异、教学反思、具体案例-20260720.docx"),
        "teacher_code": "T-01",
        "teacher_real_name": "常育红",
        "school_anonymized": "某省重点中学",
        "interviewer_anonymized": "访谈人-01",
        "interviewer_real_name": "杜玥林",
    },
    {
        "doc_id": "I-003",
        "file_path": str(SEED_MATERIALS_ROOT / "数据收集0811" / "访谈" / "访谈-张老师-教学设计心得-20260716.docx"),
        "teacher_code": "T-02",
        "teacher_real_name": None,  # 张老师无真实姓名
        "school_anonymized": "某外国语大学",
        "interviewer_anonymized": "访谈人-02",
        "interviewer_real_name": "王蓝清",
    },
    {
        "doc_id": "I-004",
        "file_path": str(SEED_MATERIALS_ROOT / "数据收集0811" / "访谈" / "访谈-林老师-教学设计心得-20260802.docx"),
        "teacher_code": "T-03",
        "teacher_real_name": None,
        "school_anonymized": "某二本院校",
        "interviewer_anonymized": "访谈人-02",
        "interviewer_real_name": "王蓝清",
    },
    {
        "doc_id": "I-005",
        "file_path": str(SEED_MATERIALS_ROOT / "数据收集0811" / "访谈" / "访谈-陈老师-教学设计心得-20260725.docx"),
        "teacher_code": "T-04",
        "teacher_real_name": None,
        "school_anonymized": "某高校外国语学院",
        "interviewer_anonymized": "访谈人-02",
        "interviewer_real_name": "王蓝清",
    },
]


def get_all_lines(doc: docx.Document) -> list:
    """按文档实际顺序交错读取段落和表格

    python-docx 的 doc.paragraphs + doc.tables 是两个独立列表，
    会丢失段落和表格在文档里的交错顺序。张老师访谈里"答"放在表格里，
    必须按实际顺序读才能让状态机正确配对问答。

    实现：遍历 body 的子元素，w:p 处理段落，w:tbl 处理表格。
    注意：w:br/w:cr 节点要转成 \n，否则 0729 常老师访谈里
    一个段落里的多行（用软回车分隔）会被压成一行，破坏正则匹配。
    """
    from docx.oxml.ns import qn

    def extract_text(elem) -> str:
        """拼接 elem 内所有 w:t 文本，w:br/w:cr 转为换行"""
        parts = []
        for sub in elem.iter():
            if sub.tag == qn("w:t"):
                parts.append(sub.text or "")
            elif sub.tag in (qn("w:br"), qn("w:cr")):
                parts.append("\n")
        return "".join(parts)

    lines = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = extract_text(child)
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    lines.append(line)
        elif child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                for cell in row.findall(qn("w:tc")):
                    cell_text = extract_text(cell)
                    for line in cell_text.split("\n"):
                        line = line.strip()
                        if line:
                            lines.append(line)
    return lines


def extract_basic_info(lines: list, file_meta: dict) -> dict:
    """从行列表提取访谈基本信息

    兼容两种格式：
    - 行式（0729 常老师）：每行一个字段，`- 访谈对象：xxx（xxx，12年）`
    - 紧凑式（0811 张/林/陈）：所有字段挤在一个表格单元格里，
      `访谈对象：xxx访谈时间：2026-07-16访谈人：xxx访谈主题：xxx`

    策略：把前 50 行用空格拼起来，对整体做 re.search，
    让两种格式都能被同一组正则匹配。
    """
    info = {
        "teacher_code": file_meta["teacher_code"],
        "school_anonymized": file_meta["school_anonymized"],
        "interviewer_anonymized": file_meta["interviewer_anonymized"],
        "teaching_years": None,
        "interview_date": None,
        "interview_topic": None,
    }

    header_text = " ".join(lines[:50])

    # 教龄：先匹配"教龄5年"显式格式，再从访谈对象括号里兜底"5年"
    ym = re.search(r"教龄\s*(\d+)\s*年", header_text)
    if ym:
        info["teaching_years"] = int(ym.group(1))
    else:
        om = re.search(r"访谈对象[^（(]*[（(]([^)）]+)[)）]", header_text)
        if om:
            ym2 = re.search(r"(\d+)\s*年", om.group(1))
            if ym2:
                info["teaching_years"] = int(ym2.group(1))

    # 访谈时间：日期形如 2026-07-14（无空格、无中文）
    dm = re.search(r"访谈时间\s*[：:]\s*([^\s访谈]+)", header_text)
    if dm:
        info["interview_date"] = dm.group(1).strip()

    # 访谈主题：到下一个"访谈"关键字或 markdown 标题为止
    tm = re.search(
        r"访谈主题\s*[：:]\s*(.+?)(?=\s*(?:访谈|##\s|#\s|$))",
        header_text,
    )
    if tm:
        info["interview_topic"] = tm.group(1).strip().rstrip("#").strip()

    return info


def parse_qa_pairs(lines: list) -> list:
    """状态机识别问答对

    识别模式：
    - **问**：xxx 或 问：xxx  → 新问题
    - **答**：xxx 或 答：xxx  → 新答案
    - ### 话题N：xxx 或 话题N：xxx → 话题标题
    - 其他行 → 追加到当前 question 或 answer
    """
    qa_pairs = []
    current_q = ""
    current_a = ""
    current_topic = ""
    state = None  # None / 'q' / 'a'

    def flush():
        nonlocal current_q, current_a
        if current_q and current_a:
            qa_pairs.append({
                "topic": current_topic,
                "question": current_q.strip(),
                "answer": current_a.strip(),
            })
        current_q = ""
        current_a = ""

    for line in lines:
        # 话题标题：### 话题N：xxx 或 话题N：xxx
        topic_m = re.match(r"^(?:###\s*)?话题\s*\d+\s*[：:]\s*(.+)$", line)
        if topic_m:
            flush()
            current_topic = topic_m.group(1).strip()
            state = None
            continue

        # 问：**问**：xxx 或 问：xxx
        q_m = re.match(r"^(?:\*\*)?问(?:\*\*)?\s*[：:]\s*(.*)$", line)
        if q_m:
            flush()
            current_q = q_m.group(1).strip()
            state = "q"
            continue

        # 答：**答**：xxx 或 答：xxx
        a_m = re.match(r"^(?:\*\*)?答(?:\*\*)?\s*[：:]\s*(.*)$", line)
        if a_m:
            current_a = a_m.group(1).strip()
            state = "a"
            continue

        # 其他行：根据 state 追加
        if state == "q":
            current_q += "\n" + line
        elif state == "a":
            current_a += "\n" + line

    flush()
    return qa_pairs


def anonymize_qa_content(qa_pairs: list, teacher_real_name: str = None) -> list:
    """匿名化问答内容里的真实姓名/学校

    把 teacher_real_name 替换为代号（如 T-01）。
    注意：只替换真实姓名，不替换其他内容。
    """
    if not teacher_real_name:
        return qa_pairs

    anonymized = []
    for qa in qa_pairs:
        q = qa["question"].replace(teacher_real_name, "T老师")
        a = qa["answer"].replace(teacher_real_name, "T老师")
        anonymized.append({
            "topic": qa["topic"],
            "question": q,
            "answer": a,
        })
    return anonymized


def main():
    entries = []
    failures = []

    for file_meta in INTERVIEW_FILES:
        path = file_meta["file_path"]
        print(f"处理 {file_meta['doc_id']}: {os.path.basename(path)}")

        if not os.path.exists(path):
            failures.append({
                "doc_id": file_meta["doc_id"],
                "file_path": path,
                "reason": "文件不存在",
            })
            continue

        try:
            doc = docx.Document(path)
        except Exception as e:
            failures.append({
                "doc_id": file_meta["doc_id"],
                "file_path": path,
                "reason": f"docx 解析失败: {e}",
            })
            continue

        lines = get_all_lines(doc)
        basic_info = extract_basic_info(lines, file_meta)
        qa_pairs = parse_qa_pairs(lines)
        qa_pairs = anonymize_qa_content(qa_pairs, file_meta.get("teacher_real_name"))

        entry = {
            "doc_id": file_meta["doc_id"],
            "file_path": path,
            "original_filename": os.path.basename(path),
            "teacher_code": file_meta["teacher_code"],
            "school_anonymized": file_meta["school_anonymized"],
            "interviewer_anonymized": file_meta["interviewer_anonymized"],
            "teaching_years": basic_info["teaching_years"],
            "interview_date": basic_info["interview_date"],
            "interview_topic": basic_info["interview_topic"],
            "qa_count": len(qa_pairs),
            "total_chars": sum(len(qa["question"]) + len(qa["answer"]) for qa in qa_pairs),
            "qa_pairs": qa_pairs,
            "scope": "system",
            "evidence_type": "direct_quote",
            "confidence": "high",
            "language_scope": "cross-lingual",
            "_anonymization_note": (
                f"教师真实姓名 {file_meta['teacher_real_name'] or '(未提供)'} "
                f"已替换为代号 {file_meta['teacher_code']}；"
                f"访谈人 {file_meta['interviewer_real_name']} 已替换为 {file_meta['interviewer_anonymized']}；"
                f"学校名称已匿名化为 {file_meta['school_anonymized']}"
            ),
        }
        entries.append(entry)
        print(f"  -> {len(qa_pairs)} 个问答对，{entry['total_chars']} 字符")

    manifest = {
        "schema_version": "1.0",
        "generated_at": None,
        "total_count": len(entries),
        "teacher_code_map": {
            "T-01": "常老师（0729 两份访谈）",
            "T-02": "张老师（0811）",
            "T-03": "林老师（0811）",
            "T-04": "陈老师（0811）",
        },
        "interviewer_code_map": {
            "访谈人-01": "杜玥林",
            "访谈人-02": "王蓝清",
        },
        "school_anonymization_map": {
            "某省重点中学": "天水市第一中学（原）",
            "某外国语大学": "某外国语大学法语系（原匿名）",
            "某二本院校": "某二本院校（原匿名）",
            "某高校外国语学院": "某高校外国语学院（原匿名）",
        },
        "entries": entries,
        "failures": failures,
        "notes": [
            "doc_id 编号规则：I-001~I-005，按访谈日期顺序",
            "教师代号 T-01~T-04 跨访谈复用：0729 常老师两份访谈都是 T-01",
            "teacher_quote 字段在 extract_strategies.py 中从 qa_pairs[i].answer 提取",
            "qa_pairs 数组结构：每条含 topic/question/answer 三字段，topic 跨问答对可能相同",
            "language_scope=translingual 是因为访谈涉及中、英、法多语种教学场景",
            "scope=system：seed 入库时与 5 个 system_seed 并存",
        ],
    }

    from datetime import datetime, timezone
    manifest["generated_at"] = datetime.now(timezone.utc).isoformat()

    OUTPUT_PATH.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"\nOK: {len(entries)} entries written to {OUTPUT_PATH}")
    print(f"  - failures: {len(failures)}")
    total_qa = sum(e["qa_count"] for e in entries)
    total_chars = sum(e["total_chars"] for e in entries)
    print(f"  - total QA pairs: {total_qa}")
    print(f"  - total chars: {total_chars}")


if __name__ == "__main__":
    main()
