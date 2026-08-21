"""
生成 lesson.json 教案元数据清单

数据来源：seed-materials/数据收集0811/教案/ 下的 24 份教案
按模板聚类分为 4 类：
- 常老师自定义（5 .docx）：英语听说/精读 1-3，自带元信息段
- 表格型（6 .doc/.docx）：郑老师大学英语 综合U1/U3、听力U1/U3/U4/U7
- 出版社配套（7 .txt GBK）：许老师 新航标职业英语·综合英语 1 第 1-6/8 单元
- 外研社大赛（6 .pdf）：郑老师法语读写 unit2/4/6/7/8 + 钱老师汉俄翻译 unit1

输出：backend/scripts/manifests/lesson.json

匿名化规则：
- 常老师 → T-01（与访谈 T-01 同一人，跨素材复用）
- 郑老师 → T-05
- 许老师 → T-06
- 钱老师 → T-07
- 学校具体名称 → "某省重点中学" / "某高校外国语学院" / "某高职院校"
"""

import json
import os
import re
import sys
from pathlib import Path
from typing import Tuple

SEED_MATERIALS_ROOT = Path(r"C:\Users\ht\Documents\outeye3.0\seed-materials")
LESSON_ROOT = SEED_MATERIALS_ROOT / "数据收集0811" / "教案"
OUTPUT_PATH = Path(__file__).parent / "lesson.json"


LESSON_FILES = [
    # === 常老师自定义（5 .docx）===
    {
        "doc_id": "L-001",
        "file_path": str(LESSON_ROOT / "教案-英语听说1-课文教学-常老师-2025.docx"),
        "teacher_code": "T-01",
        "school_anonymized": "某省重点中学",
        "template_cluster": "常老师自定义",
    },
    {
        "doc_id": "L-002",
        "file_path": str(LESSON_ROOT / "教案-英语听说2-课文教学-常老师-2025.docx"),
        "teacher_code": "T-01",
        "school_anonymized": "某省重点中学",
        "template_cluster": "常老师自定义",
    },
    {
        "doc_id": "L-003",
        "file_path": str(LESSON_ROOT / "教案-英语精读1-课文教学-常老师-2025.docx"),
        "teacher_code": "T-01",
        "school_anonymized": "某省重点中学",
        "template_cluster": "常老师自定义",
    },
    {
        "doc_id": "L-004",
        "file_path": str(LESSON_ROOT / "教案-英语精读2-课文教学-常老师-2025.docx"),
        "teacher_code": "T-01",
        "school_anonymized": "某省重点中学",
        "template_cluster": "常老师自定义",
    },
    {
        "doc_id": "L-005",
        "file_path": str(LESSON_ROOT / "教案-英语精读3-课文教学-常老师-2025.docx"),
        "teacher_code": "T-01",
        "school_anonymized": "某省重点中学",
        "template_cluster": "常老师自定义",
    },
    # === 表格型（6 .doc/.docx）郑老师大学英语 ===
    {
        "doc_id": "L-006",
        "file_path": str(LESSON_ROOT / "教案-大学英语-综合-U1-郑老师-2023.doc"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    {
        "doc_id": "L-007",
        "file_path": str(LESSON_ROOT / "教案-大学英语听力-U1-郑老师-2023.doc"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    {
        "doc_id": "L-008",
        "file_path": str(LESSON_ROOT / "教案-大学英语听力-U3-郑老师-2023.doc"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    {
        "doc_id": "L-009",
        "file_path": str(LESSON_ROOT / "教案-大学英语听力-U4-郑老师-2023.doc"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    {
        "doc_id": "L-010",
        "file_path": str(LESSON_ROOT / "教案-大学英语听力-U7-郑老师-2023.doc"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    {
        "doc_id": "L-011",
        "file_path": str(LESSON_ROOT / "教案-大学英语综合-U3-郑老师-2023.docx"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "表格型",
    },
    # === 出版社配套（7 .txt GBK）许老师 ===
    {
        "doc_id": "L-012",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第1单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-013",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第2单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-014",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第3单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-015",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第4单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-016",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第5单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-017",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第6单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    {
        "doc_id": "L-018",
        "file_path": str(LESSON_ROOT / "教案-新航标职业英语·综合英语1-第8单元-许老师-2013.txt"),
        "teacher_code": "T-06",
        "school_anonymized": "某高职院校",
        "template_cluster": "出版社配套",
    },
    # === 外研社大赛（6 .pdf）===
    {
        "doc_id": "L-019",
        "file_path": str(LESSON_ROOT / "教案-法语读写-综合课unit2-郑老师-2022.pdf"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
    {
        "doc_id": "L-020",
        "file_path": str(LESSON_ROOT / "教案-法语读写-综合课unit4-郑老师-2022.pdf"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
    {
        "doc_id": "L-021",
        "file_path": str(LESSON_ROOT / "教案-法语读写-综合课unit6-郑老师-2022.pdf"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
    {
        "doc_id": "L-022",
        "file_path": str(LESSON_ROOT / "教案-法语读写-综合课unit7-郑老师-2022.pdf"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
    {
        "doc_id": "L-023",
        "file_path": str(LESSON_ROOT / "教案-法语读写-综合课unit8-郑老师-2022.pdf"),
        "teacher_code": "T-05",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
    {
        "doc_id": "L-024",
        "file_path": str(LESSON_ROOT / "教案-汉俄翻译-翻译教学unit1-钱老师-2022.pdf"),
        "teacher_code": "T-07",
        "school_anonymized": "某高校外国语学院",
        "template_cluster": "外研社大赛",
    },
]


def parse_filename_metadata(filename: str) -> dict:
    """从文件名解析元数据

    文件名格式：教案-{course+unit+lesson_type}-{teacher}-{year}.{ext}
    例：教案-英语听说1-课文教学-常老师-2025.docx
       -> course="英语听说1", lesson_type="课文教学", teacher="常老师", year=2025
    """
    m = re.match(
        r"^教案-(.+)-([^\-]+老师)-(\d{4})\.(.+)$",
        filename,
    )
    if not m:
        return {
            "course": None,
            "unit": None,
            "lesson_type": None,
            "teacher": None,
            "year": None,
            "file_format": None,
        }

    middle = m.group(1)  # course+unit+lesson_type 组合
    teacher = m.group(2)
    year = int(m.group(3))
    ext = "." + m.group(4)

    # 从 middle 中分离 unit 和 lesson_type
    # 常见模式：
    # - "英语听说1-课文教学"：course=英语听说1, lesson_type=课文教学
    # - "新航标职业英语·综合英语1-第1单元"：course=新航标职业英语·综合英语1, unit=第1单元
    # - "大学英语-综合-U1"：course=大学英语-综合, unit=U1
    # - "法语读写-综合课unit2"：course=法语读写, lesson_type=综合课unit2
    # - "汉俄翻译-翻译教学unit1"：course=汉俄翻译, lesson_type=翻译教学unit1

    unit = None
    lesson_type = None
    course = middle

    # 提取 unit（如 "第1单元" 或 "U1"）
    unit_m = re.search(r"(第\d+单元|[Uu]\d+)", middle)
    if unit_m:
        unit = unit_m.group(1)
        # 把 unit 从 middle 中切掉，剩下的就是 course+lesson_type
        before_unit = middle[:unit_m.start()].rstrip("-")
        after_unit = middle[unit_m.end():].lstrip("-")
        if before_unit and after_unit:
            course = before_unit
            lesson_type = after_unit
        elif before_unit:
            course = before_unit
            lesson_type = None
        elif after_unit:
            course = middle
            lesson_type = None
    else:
        # 没有 unit，按 "-" 切分 course 和 lesson_type
        parts = middle.rsplit("-", 1)
        if len(parts) == 2:
            course = parts[0]
            lesson_type = parts[1]
        else:
            course = middle
            lesson_type = None

    # 如果 lesson_type 包含 "unit" 关键字（如 "综合课unit2"），尝试提取 unit
    if lesson_type and not unit:
        unit_m2 = re.search(r"unit\s*(\d+)", lesson_type, re.IGNORECASE)
        if unit_m2:
            unit = f"unit{unit_m2.group(1)}"
            # 从 lesson_type 里去掉 unit 部分，避免和 unit 字段重复
            lesson_type = re.sub(
                r"\s*unit\s*\d+", "", lesson_type, flags=re.IGNORECASE
            ).strip() or None

    return {
        "course": course,
        "unit": unit,
        "lesson_type": lesson_type,
        "teacher": teacher,
        "year": year,
        "file_format": ext,
    }


def read_docx(path: str) -> str:
    """读 .docx 全文"""
    import docx
    doc = docx.Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 表格内容也加上
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def read_doc_win32(path: str, max_retries: int = 3) -> str:
    """用 win32com 读 .doc 文件（仅 Windows）

    Word COM 偶发 RPC 服务器不可用 / 卡死，Open/Close/Quit 任何一步失败都不影响已读到的文本。
    失败时重试最多 max_retries 次，每次重试前清理残留 Word 进程。
    """
    import time
    import win32com.client as wc
    import pythoncom

    abs_path = os.path.abspath(path)

    last_err = None
    for attempt in range(max_retries):
        pythoncom.CoInitialize()
        word = None
        doc = None
        text = ""
        try:
            word = wc.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(abs_path, ReadOnly=True)
            text = doc.Content.Text
            return text
        except Exception as e:
            last_err = e
            # 重试前等一会儿，让 Word 释放资源
            time.sleep(1.0 * (attempt + 1))
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    raise RuntimeError(f"Word COM 读取失败（重试 {max_retries} 次）: {last_err}")


def read_txt_with_encoding(path: str) -> str:
    """读 .txt，尝试 UTF-8、GBK、GB18030"""
    for enc in ["utf-8", "gbk", "gb18030", "utf-16"]:
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # 最后兜底
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def read_pdf_pypdf2(path: str) -> str:
    """用 PyPDF2 读 PDF（CJK 兼容性差，服务器上会用 PyMuPDF）"""
    import PyPDF2
    text_parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def read_file_content(path: str) -> Tuple[str, str]:
    """根据扩展名读文件，返回 (text, reader_used)

    读后统一清理控制字符：\x07（BELL）等会卡在词与数字之间，
    让 "教学时数\\x076" 这种文本无法被正则匹配。
    """
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        text, reader = read_docx(path), "python-docx"
    elif suffix == ".doc":
        text, reader = read_doc_win32(path), "win32com"
    elif suffix == ".txt":
        text, reader = read_txt_with_encoding(path), "txt-encoding"
    elif suffix == ".pdf":
        text, reader = read_pdf_pypdf2(path), "PyPDF2"
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")

    # 清理控制字符（BELL \x07、垂直制表 \x0B、换页 \x0C 等），保留 \n\r\t
    text = "".join(
        c for c in text
        if c == "\n" or c == "\r" or c == "\t"
        or (ord(c) >= 0x20 and ord(c) != 0x07)
    )
    return text, reader


def extract_content_metadata(text: str, template_cluster: str) -> dict:
    """从全文提取 student_level / score 等内容元数据

    不同模板的元数据位置：
    - 常老师自定义：开头段 `学生水平：[A2]` `课时：[1课时]`（值带方括号）
    - 表格型：开头 `教学时数 6`
    - 出版社配套：开头结构不固定，正文多为英文
    - 外研社大赛：开头 `授课对象 法语专业三年级` `课程名称 高级法语`
    """
    info = {
        "student_level": None,
        "class_hours": None,
        "score": None,
    }

    # 学生水平：[A2] 或 A2 或 授课对象：法语专业三年级
    m = re.search(r"学生水平\s*[：:]\s*\[?([A-C]\d)\]?", text)
    if m:
        info["student_level"] = m.group(1)
    if not info["student_level"]:
        m = re.search(r"授课对象\s*[：:]?\s*(.{0,30}?(?:一年级|二年级|三年级|四年级))", text)
        if m:
            info["student_level"] = m.group(1).strip()
    if not info["student_level"]:
        m = re.search(r"(英语|法语|俄语|日语|德语|西语)专业\s*(一|二|三|四)年级", text)
        if m:
            info["student_level"] = m.group(0).strip()

    # 课时：[1课时] / 1课时 / 6学时 / 教学时数 6
    m = re.search(r"课时\s*[：:]\s*\[?(\d+)\s*课时?\]?", text)
    if m:
        info["class_hours"] = int(m.group(1))
    if not info["class_hours"]:
        m = re.search(r"教学时数\s*[：:]?\s*(\d+)", text)
        if m:
            info["class_hours"] = int(m.group(1))
    if not info["class_hours"]:
        m = re.search(r"(\d+)\s*学时", text)
        if m:
            info["class_hours"] = int(m.group(1))

    # 评分：教案若含专家评分分项可提取
    m = re.search(r"评分\s*[：:]\s*(\d+(?:\.\d+)?)", text)
    if m:
        info["score"] = float(m.group(1))

    return info


def main():
    entries = []
    failures = []

    for file_meta in LESSON_FILES:
        path = file_meta["file_path"]
        filename = os.path.basename(path)
        print(f"处理 {file_meta['doc_id']}: {filename}")

        if not os.path.exists(path):
            failures.append({
                "doc_id": file_meta["doc_id"],
                "file_path": path,
                "reason": "文件不存在",
            })
            continue

        # 从文件名提取元数据
        fn_meta = parse_filename_metadata(filename)

        # 读全文
        try:
            text, reader_used = read_file_content(path)
        except Exception as e:
            failures.append({
                "doc_id": file_meta["doc_id"],
                "file_path": path,
                "reason": f"读取失败: {e}",
            })
            continue

        # 从内容提取 student_level / class_hours / score
        content_meta = extract_content_metadata(text, file_meta["template_cluster"])

        # PDF PyPDF2 在本地对 GBK-EUC-H 编码有 mojibake，标记为待服务器验证
        is_pdf_mojibake = (
            file_meta["file_format"] if isinstance(file_meta.get("file_format"), str) else Path(path).suffix.lower()
        ) == ".pdf" or Path(path).suffix.lower() == ".pdf"

        _pending_verification = False
        _verification_note = None
        if is_pdf_mojibake and reader_used == "PyPDF2":
            _pending_verification = True
            _verification_note = (
                "本地用 PyPDF2 解析存在 GBK-EUC-H 编码 mojibake，"
                "服务器用 PyMuPDF 重新解析后内容应正常"
            )

        entry = {
            "doc_id": file_meta["doc_id"],
            "file_path": path,
            "original_filename": filename,
            "teacher_code": file_meta["teacher_code"],
            "school_anonymized": file_meta["school_anonymized"],
            "template_cluster": file_meta["template_cluster"],
            "course": fn_meta["course"],
            "unit": fn_meta["unit"],
            "lesson_type": fn_meta["lesson_type"],
            "year": fn_meta["year"],
            "file_format": fn_meta["file_format"],
            "student_level": content_meta["student_level"],
            "class_hours": content_meta["class_hours"],
            "score": content_meta["score"],
            "total_chars": len(text),
            "text_preview": text[:200].replace("\n", " ").strip(),
            "reader_used": reader_used,
            "_pending_verification": _pending_verification,
            "_verification_note": _verification_note,
            "_anonymization_note": (
                f"教师代号 {file_meta['teacher_code']}（教案中署名为 "
                f"{fn_meta['teacher'] or '未知姓+老师'}）；"
                f"学校名称已匿名化为 {file_meta['school_anonymized']}"
            ),
        }
        entries.append(entry)
        print(f"  -> {entry['total_chars']} chars, "
              f"level={entry['student_level']}, "
              f"hours={entry['class_hours']}, "
              f"score={entry['score']}")

    manifest = {
        "schema_version": "1.0",
        "generated_at": None,
        "total_count": len(entries),
        "teacher_code_map": {
            "T-01": "常老师（与访谈 T-01 跨素材复用）",
            "T-05": "郑老师（教案 6 份：大学英语 5 + 法语读写 5）",
            "T-06": "许老师（新航标职业英语 7 单元）",
            "T-07": "钱老师（汉俄翻译 1 份）",
        },
        "school_anonymization_map": {
            "某省重点中学": "天水市第一中学（原）",
            "某高校外国语学院": "某高校外国语学院（原匿名）",
            "某高职院校": "某高职院校（原匿名）",
        },
        "template_cluster_map": {
            "常老师自定义": "5 .docx，英语听说/精读，开头有元信息段",
            "表格型": "6 .doc/.docx，郑老师大学英语，主体为表格",
            "出版社配套": "7 .txt GBK，许老师新航标职业英语，正文为主",
            "外研社大赛": "6 .pdf，郑老师法语读写+钱老师汉俄翻译",
        },
        "entries": entries,
        "failures": failures,
        "notes": [
            "doc_id 编号规则：L-001~L-024，按模板聚类+教师+年份顺序",
            "T-01 与访谈清单的 T-01 是同一教师，跨素材复用",
            "学生水平提取规则：CEFR 等级（A1-C2）或 专业年级（如 法语专业三年级）",
            "PDF 在本地用 PyPDF2 解析有 mojibake，标记 _pending_verification，"
            "服务器用 PyMuPDF 重新解析后内容应正常",
            "score 字段目前均提取为 None，教案无显式评分；如需专家评分可后续在 review 阶段补",
            "scope=system：seed 入库时与现有 5 个 system_seed 并存",
        ],
    }

    from datetime import datetime, timezone
    manifest["generated_at"] = datetime.now(timezone.utc).isoformat()

    OUTPUT_PATH.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"\nOK: {len(entries)} entries written to {OUTPUT_PATH}")
    print(f"  - failures: {len(failures)}")
    total_chars = sum(e["total_chars"] for e in entries)
    print(f"  - total chars: {total_chars}")
    pending = sum(1 for e in entries if e.get("_pending_verification"))
    print(f"  - pending verification (PDF mojibake): {pending}")


if __name__ == "__main__":
    main()
